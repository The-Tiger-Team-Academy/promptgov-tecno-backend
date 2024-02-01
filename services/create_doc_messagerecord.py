from docx import Document
from fastapi import APIRouter, HTTPException
from entity.messagerecord import CreateDocMessageRecord
from google.cloud import storage
import subprocess
import os

# Initialize Google Cloud Storage Client
storage_client = storage.Client.from_service_account_json("storage-admin.json")
bucket_name = "storagepromtgovfile"

router = APIRouter(
    tags=["Create Services"],
    responses={404: {"description": "Not Found"}},
)

def convert_to_pdf_libreoffice(filepath):
    pdfpath = filepath.replace(".docx", ".pdf")
    try:
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf', 
            '--outdir', os.path.dirname(filepath), filepath
        ], check=True)
        if not os.path.exists(pdfpath):
            raise Exception("PDF conversion failed: Output file not created.")
        return pdfpath
    except subprocess.CalledProcessError as e:
        raise HTTPException(status_code=500, detail=f"Error in PDF conversion: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def upload_to_storage(filepath, bucket_name):
    bucket = storage_client.get_bucket(bucket_name)
    blob = bucket.blob(os.path.basename(filepath))
    blob.upload_from_filename(filepath)
    return blob.public_url

@router.post("/create_doc/messageRecord/")
async def create_doc_messageRecord(doc_data: CreateDocMessageRecord):
    try:
        # Replace content in the document
        doc = Document("./templatedoc/messageRecordTemplate.docx")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for old, new in doc_data.dict().items():
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

        # Save the modified document
        filepath = f"{doc_data.STORY}.docx"
        doc.save(filepath)

        # Convert to PDF
        pdfpath = convert_to_pdf_libreoffice(filepath)

        # Upload to Google Cloud Storage
        docx_public_url = upload_to_storage(filepath, bucket_name)
        pdf_public_url = upload_to_storage(pdfpath, bucket_name)

        # Cleanup local files
        os.remove(pdfpath)
        os.remove(filepath)

        return {"url": docx_public_url, "url_pdf": pdf_public_url}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

